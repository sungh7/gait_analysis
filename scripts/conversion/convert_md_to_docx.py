import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def parse_markdown_to_docx(md_path, docx_path):
    doc = docx.Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_mode = False
    table_data = []
    
    # Regex patterns
    img_pattern = re.compile(r'!\[(.*?)\]\((.*?)\)')
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')
    italic_pattern = re.compile(r'\*(.*?)\*')

    for line in lines:
        line = line.strip()
        
        # Skip empty lines unless ending a table
        if not line:
            if table_mode:
                create_table(doc, table_data)
                table_mode = False
                table_data = []
            continue

        # 1. Tables
        if line.startswith('|'):
            table_mode = True
            # Parse row
            cells = [c.strip() for c in line.split('|') if c]
            # Skip separator line (e.g., |---|---|)
            if '---' in line:
                continue
            table_data.append(cells)
            continue
        
        # If we were in table mode but line doesn't start with |, table ended
        if table_mode:
            create_table(doc, table_data)
            table_mode = False
            table_data = []

        # 2. Images
        img_match = img_pattern.match(line)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add caption
                    caption_para = doc.add_paragraph(caption)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_para.style = 'Caption'
                except Exception as e:
                    doc.add_paragraph(f"[Image Error: {e}]")
            else:
                doc.add_paragraph(f"[Image not found: {img_path}]")
            continue

        # 3. Headers
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=min(level, 9))
            continue

        # 4. Lists
        if line.startswith('- '):
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            format_text(p, text, bold_pattern, italic_pattern)
            continue

        # 5. Normal Paragraph
        p = doc.add_paragraph()
        format_text(p, line, bold_pattern, italic_pattern)

    # End of file check for table
    if table_mode:
        create_table(doc, table_data)

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

def format_text(paragraph, text, bold_pattern, italic_pattern):
    # This is a simplified formatter. It doesn't handle nested bold/italic well.
    # We'll split by bold first.
    
    # A better approach for mixed formatting is hard without a full parser.
    # We will try a simple approach: 
    # If the whole line is bold, make it bold.
    # Otherwise, just add as text. 
    # Implementing full inline formatting parsing is complex for a script.
    
    # Let's try to handle bold at least.
    parts = bold_pattern.split(text)
    # If match, parts will be [pre, bold_text, post, bold_text, post...]
    
    # Check if we actually matched
    if len(parts) > 1:
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            if i % 2 == 1: # Odd indices are the captured groups (bold text)
                run.bold = True
    else:
        paragraph.add_run(text)

def create_table(doc, data):
    if not data:
        return
    
    rows = len(data)
    cols = len(data[0])
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                row.cells[j].text = cell_text

if __name__ == "__main__":
    md_file = "/data/gait/PAPER_MediaPipe_Individual_Gait_Analysis.md"
    docx_file = "/data/gait/PAPER_MediaPipe_Individual_Gait_Analysis.docx"
    parse_markdown_to_docx(md_file, docx_file)
